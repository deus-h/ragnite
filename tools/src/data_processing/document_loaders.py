"""
Document Loaders

This module provides various document loaders for different file formats.
"""

import os
import re
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Union, Iterator
import json
import csv
from pathlib import Path
import logging

# Optional imports for specific document types
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pptx
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


# Set up logging
logger = logging.getLogger(__name__)


class Document:
    """
    A simple document class to store text content and metadata.
    """
    
    def __init__(
        self, 
        content: str, 
        metadata: Optional[Dict[str, Any]] = None
    ):
        """
        Initialize a document with content and optional metadata.
        
        Args:
            content: The text content of the document
            metadata: Optional metadata associated with the document
        """
        self.content = content
        self.metadata = metadata or {}
    
    def __repr__(self) -> str:
        """String representation of the document."""
        return f"Document(content_length={len(self.content)}, metadata={self.metadata})"


class BaseLoader(ABC):
    """
    Abstract base class for document loaders.
    """
    
    @abstractmethod
    def load(self) -> List[Document]:
        """
        Load document(s) from a source.
        
        Returns:
            List of Document objects
        """
        pass
    
    @abstractmethod
    def lazy_load(self) -> Iterator[Document]:
        """
        Lazily load document(s) from a source.
        
        Returns:
            Iterator of Document objects
        """
        pass


class TextLoader(BaseLoader):
    """
    Loader for plain text files.
    """
    
    def __init__(
        self, 
        file_path: str, 
        encoding: str = "utf-8", 
        autodetect_encoding: bool = False
    ):
        """
        Initialize a text file loader.
        
        Args:
            file_path: Path to the text file
            encoding: Encoding of the text file
            autodetect_encoding: Whether to try to autodetect the encoding
        """
        self.file_path = file_path
        self.encoding = encoding
        self.autodetect_encoding = autodetect_encoding
    
    def _get_encoding(self) -> str:
        """
        Get the encoding to use, possibly autodetected.
        
        Returns:
            Encoding to use
        """
        if not self.autodetect_encoding:
            return self.encoding
        
        # Try to autodetect encoding
        try:
            import chardet
            with open(self.file_path, "rb") as f:
                result = chardet.detect(f.read())
            detected_encoding = result["encoding"]
            confidence = result["confidence"]
            
            if detected_encoding and confidence > 0.7:
                logger.info(f"Detected encoding: {detected_encoding} with confidence {confidence}")
                return detected_encoding
        except ImportError:
            logger.warning("chardet not available for encoding autodetection")
        except Exception as e:
            logger.warning(f"Error autodetecting encoding: {e}")
        
        return self.encoding
    
    def load(self) -> List[Document]:
        """
        Load a text file.
        
        Returns:
            List containing a single Document with the file content
        """
        try:
            encoding = self._get_encoding()
            with open(self.file_path, "r", encoding=encoding) as f:
                text = f.read()
            
            metadata = {
                "source": self.file_path,
                "file_type": "text",
                "file_name": os.path.basename(self.file_path),
                "file_path": os.path.abspath(self.file_path),
                "file_size": os.path.getsize(self.file_path),
                "encoding": encoding,
            }
            
            return [Document(content=text, metadata=metadata)]
        except Exception as e:
            logger.error(f"Error loading text file {self.file_path}: {e}")
            raise
    
    def lazy_load(self) -> Iterator[Document]:
        """
        Lazily load a text file.
        
        Yields:
            Document object with the file content
        """
        try:
            encoding = self._get_encoding()
            with open(self.file_path, "r", encoding=encoding) as f:
                text = f.read()
            
            metadata = {
                "source": self.file_path,
                "file_type": "text",
                "file_name": os.path.basename(self.file_path),
                "file_path": os.path.abspath(self.file_path),
                "file_size": os.path.getsize(self.file_path),
                "encoding": encoding,
            }
            
            yield Document(content=text, metadata=metadata)
        except Exception as e:
            logger.error(f"Error lazy loading text file {self.file_path}: {e}")
            raise


class PDFLoader(BaseLoader):
    """
    Loader for PDF files using PyMuPDF.
    """
    
    def __init__(
        self, 
        file_path: str, 
        extract_images: bool = False,
        extract_tables: bool = False,
        include_page_numbers: bool = True,
        password: Optional[str] = None
    ):
        """
        Initialize a PDF loader.
        
        Args:
            file_path: Path to the PDF file
            extract_images: Whether to extract images (as metadata)
            extract_tables: Whether to extract tables
            include_page_numbers: Whether to include page numbers in metadata
            password: Password for encrypted PDFs
        """
        self.file_path = file_path
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.include_page_numbers = include_page_numbers
        self.password = password
        
        if not PYMUPDF_AVAILABLE:
            raise ImportError(
                "PyMuPDF (fitz) is required for PDFLoader. "
                "Install it with `pip install pymupdf`."
            )
    
    def load(self) -> List[Document]:
        """
        Load a PDF file, one Document per page.
        
        Returns:
            List of Document objects, one per page
        """
        try:
            doc = fitz.open(self.file_path, password=self.password)
            documents = []
            
            # Extract PDF metadata
            pdf_metadata = doc.metadata
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                
                # Create page metadata
                metadata = {
                    "source": self.file_path,
                    "file_type": "pdf",
                    "file_name": os.path.basename(self.file_path),
                    "file_path": os.path.abspath(self.file_path),
                    "file_size": os.path.getsize(self.file_path),
                    "total_pages": len(doc),
                    "pdf_metadata": pdf_metadata,
                }
                
                if self.include_page_numbers:
                    metadata["page_number"] = page_num + 1
                
                # Extract images if requested
                if self.extract_images:
                    images = []
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        if base_image:
                            image_data = {
                                "xref": xref,
                                "width": base_image["width"],
                                "height": base_image["height"],
                                "ext": base_image["ext"],
                                "image_index": img_index,
                            }
                            images.append(image_data)
                    
                    metadata["images"] = images
                
                # Extract tables if requested
                if self.extract_tables:
                    # Note: PyMuPDF doesn't have native table extraction
                    # This is a placeholder for future implementation
                    pass
                
                documents.append(Document(content=text, metadata=metadata))
            
            return documents
        except Exception as e:
            logger.error(f"Error loading PDF file {self.file_path}: {e}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()
    
    def lazy_load(self) -> Iterator[Document]:
        """
        Lazily load a PDF file, one Document per page.
        
        Yields:
            Document objects, one per page
        """
        try:
            doc = fitz.open(self.file_path, password=self.password)
            
            # Extract PDF metadata
            pdf_metadata = doc.metadata
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                
                # Create page metadata
                metadata = {
                    "source": self.file_path,
                    "file_type": "pdf",
                    "file_name": os.path.basename(self.file_path),
                    "file_path": os.path.abspath(self.file_path),
                    "file_size": os.path.getsize(self.file_path),
                    "total_pages": len(doc),
                    "pdf_metadata": pdf_metadata,
                }
                
                if self.include_page_numbers:
                    metadata["page_number"] = page_num + 1
                
                # Extract images if requested
                if self.extract_images:
                    images = []
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        if base_image:
                            image_data = {
                                "xref": xref,
                                "width": base_image["width"],
                                "height": base_image["height"],
                                "ext": base_image["ext"],
                                "image_index": img_index,
                            }
                            images.append(image_data)
                    
                    metadata["images"] = images
                
                yield Document(content=text, metadata=metadata)
        except Exception as e:
            logger.error(f"Error lazy loading PDF file {self.file_path}: {e}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()


class HTMLLoader(BaseLoader):
    """
    Loader for HTML files using BeautifulSoup.
    """
    
    def __init__(
        self, 
        file_path: Optional[str] = None,
        url: Optional[str] = None,
        html_content: Optional[str] = None,
        encoding: str = "utf-8",
        bs4_parser: str = "html.parser",
        extract_links: bool = False,
        extract_metadata: bool = True,
        include_comments: bool = False,
        strip_js: bool = True,
    ):
        """
        Initialize an HTML loader.
        
        Args:
            file_path: Path to the HTML file
            url: URL to fetch HTML content from
            html_content: HTML content as a string
            encoding: Encoding to use when reading from file
            bs4_parser: BeautifulSoup parser to use
            extract_links: Whether to extract links
            extract_metadata: Whether to extract metadata
            include_comments: Whether to include HTML comments
            strip_js: Whether to strip JavaScript
        """
        self.file_path = file_path
        self.url = url
        self.html_content = html_content
        self.encoding = encoding
        self.bs4_parser = bs4_parser
        self.extract_links = extract_links
        self.extract_metadata = extract_metadata
        self.include_comments = include_comments
        self.strip_js = strip_js
        
        if not BS4_AVAILABLE:
            raise ImportError(
                "BeautifulSoup4 is required for HTMLLoader. "
                "Install it with `pip install beautifulsoup4`."
            )
        
        if not file_path and not url and not html_content:
            raise ValueError(
                "Either file_path, url, or html_content must be provided."
            )
        
        # If URL is provided, check for requests
        if url:
            try:
                import requests
                self.requests_available = True
            except ImportError:
                raise ImportError(
                    "Requests is required for loading from URLs. "
                    "Install it with `pip install requests`."
                )
    
    def _get_html_content(self) -> str:
        """
        Get the HTML content from the source.
        
        Returns:
            HTML content as a string
        """
        if self.html_content:
            return self.html_content
        
        if self.file_path:
            with open(self.file_path, "r", encoding=self.encoding) as f:
                return f.read()
        
        if self.url:
            import requests
            response = requests.get(self.url)
            response.raise_for_status()
            return response.text
        
        raise ValueError("No HTML source available.")
    
    def _parse_html(self, html_content: str) -> BeautifulSoup:
        """
        Parse HTML content with BeautifulSoup.
        
        Args:
            html_content: HTML content as a string
            
        Returns:
            BeautifulSoup object
        """
        soup = BeautifulSoup(html_content, self.bs4_parser)
        
        # Strip JavaScript if requested
        if self.strip_js:
            for script in soup.find_all("script"):
                script.decompose()
            
            # Remove inline JavaScript
            for tag in soup.find_all(lambda tag: any(attr.startswith('on') for attr in tag.attrs)):
                for attr in list(tag.attrs):
                    if attr.startswith('on'):
                        del tag.attrs[attr]
        
        # Remove comments if requested
        if not self.include_comments:
            for comment in soup.find_all(text=lambda text: isinstance(text, BeautifulSoup.Comment)):
                comment.extract()
        
        return soup
    
    def _extract_text(self, soup: BeautifulSoup) -> str:
        """
        Extract text content from BeautifulSoup object.
        
        Args:
            soup: BeautifulSoup object
            
        Returns:
            Extracted text content
        """
        text = soup.get_text(separator=" ", strip=True)
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def _extract_metadata(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """
        Extract metadata from BeautifulSoup object.
        
        Args:
            soup: BeautifulSoup object
            
        Returns:
            Dictionary of metadata
        """
        metadata = {}
        
        # Extract title
        title_tag = soup.find("title")
        if title_tag:
            metadata["title"] = title_tag.get_text(strip=True)
        
        # Extract meta tags
        meta_tags = {}
        for tag in soup.find_all("meta"):
            name = tag.get("name") or tag.get("property")
            content = tag.get("content")
            if name and content:
                meta_tags[name] = content
        
        if meta_tags:
            metadata["meta_tags"] = meta_tags
        
        # Extract links if requested
        if self.extract_links:
            links = []
            for link in soup.find_all("a", href=True):
                links.append({
                    "text": link.get_text(strip=True),
                    "href": link["href"],
                })
            
            if links:
                metadata["links"] = links
        
        return metadata
    
    def load(self) -> List[Document]:
        """
        Load an HTML document.
        
        Returns:
            List containing a single Document with the HTML content
        """
        try:
            html_content = self._get_html_content()
            soup = self._parse_html(html_content)
            text = self._extract_text(soup)
            
            metadata = {
                "file_type": "html",
            }
            
            if self.file_path:
                metadata.update({
                    "source": self.file_path,
                    "file_name": os.path.basename(self.file_path),
                    "file_path": os.path.abspath(self.file_path),
                    "file_size": os.path.getsize(self.file_path),
                })
            
            if self.url:
                metadata["source"] = self.url
            
            if self.extract_metadata:
                metadata.update(self._extract_metadata(soup))
            
            return [Document(content=text, metadata=metadata)]
        except Exception as e:
            source = self.file_path or self.url or "html_content"
            logger.error(f"Error loading HTML from {source}: {e}")
            raise
    
    def lazy_load(self) -> Iterator[Document]:
        """
        Lazily load an HTML document.
        
        Yields:
            Document object with the HTML content
        """
        try:
            html_content = self._get_html_content()
            soup = self._parse_html(html_content)
            text = self._extract_text(soup)
            
            metadata = {
                "file_type": "html",
            }
            
            if self.file_path:
                metadata.update({
                    "source": self.file_path,
                    "file_name": os.path.basename(self.file_path),
                    "file_path": os.path.abspath(self.file_path),
                    "file_size": os.path.getsize(self.file_path),
                })
            
            if self.url:
                metadata["source"] = self.url
            
            if self.extract_metadata:
                metadata.update(self._extract_metadata(soup))
            
            yield Document(content=text, metadata=metadata)
        except Exception as e:
            source = self.file_path or self.url or "html_content"
            logger.error(f"Error lazy loading HTML from {source}: {e}")
            raise


class MarkdownLoader(BaseLoader):
    """
    Loader for Markdown files.
    """
    
    def __init__(
        self, 
        file_path: str, 
        encoding: str = "utf-8",
        convert_to_html: bool = False,
        extract_front_matter: bool = True,
    ):
        """
        Initialize a Markdown loader.
        
        Args:
            file_path: Path to the Markdown file
            encoding: Encoding of the file
            convert_to_html: Whether to convert Markdown to HTML
            extract_front_matter: Whether to extract YAML front matter
        """
        self.file_path = file_path
        self.encoding = encoding
        self.convert_to_html = convert_to_html
        self.extract_front_matter = extract_front_matter
        
        if convert_to_html and not MARKDOWN_AVAILABLE:
            raise ImportError(
                "The markdown package is required for HTML conversion. "
                "Install it with `pip install markdown`."
            )
    
    def _extract_front_matter(self, text: str) -> Tuple[Dict[str, Any], str]:
        """
        Extract YAML front matter from Markdown text.
        
        Args:
            text: Markdown text
            
        Returns:
            Tuple of (front_matter_dict, remaining_text)
        """
        front_matter = {}
        
        # Check for YAML front matter (between --- lines)
        front_matter_match = re.match(r'^---\s*\n(.*?)\n---\s*\n', text, re.DOTALL)
        if front_matter_match:
            front_matter_text = front_matter_match.group(1)
            remaining_text = text[front_matter_match.end():]
            
            # Parse YAML front matter
            try:
                import yaml
                front_matter = yaml.safe_load(front_matter_text)
                if not isinstance(front_matter, dict):
                    front_matter = {}
            except ImportError:
                logger.warning("PyYAML not available for front matter parsing")
                # Simple key-value parsing as fallback
                for line in front_matter_text.split("\n"):
                    if ":" in line:
                        key, value = line.split(":", 1)
                        front_matter[key.strip()] = value.strip()
            except Exception as e:
                logger.warning(f"Error parsing front matter: {e}")
            
            return front_matter, remaining_text
        
        return {}, text
    
    def load(self) -> List[Document]:
        """
        Load a Markdown file.
        
        Returns:
            List containing a single Document with the Markdown content
        """
        try:
            with open(self.file_path, "r", encoding=self.encoding) as f:
                text = f.read()
            
            front_matter = {}
            if self.extract_front_matter:
                front_matter, text = self._extract_front_matter(text)
            
            content = text
            if self.convert_to_html:
                content = markdown.markdown(text)
            
            metadata = {
                "source": self.file_path,
                "file_type": "markdown",
                "file_name": os.path.basename(self.file_path),
                "file_path": os.path.abspath(self.file_path),
                "file_size": os.path.getsize(self.file_path),
            }
            
            if front_matter:
                metadata["front_matter"] = front_matter
            
            return [Document(content=content, metadata=metadata)]
        except Exception as e:
            logger.error(f"Error loading Markdown file {self.file_path}: {e}")
            raise
    
    def lazy_load(self) -> Iterator[Document]:
        """
        Lazily load a Markdown file.
        
        Yields:
            Document object with the Markdown content
        """
        try:
            with open(self.file_path, "r", encoding=self.encoding) as f:
                text = f.read()
            
            front_matter = {}
            if self.extract_front_matter:
                front_matter, text = self._extract_front_matter(text)
            
            content = text
            if self.convert_to_html:
                content = markdown.markdown(text)
            
            metadata = {
                "source": self.file_path,
                "file_type": "markdown",
                "file_name": os.path.basename(self.file_path),
                "file_path": os.path.abspath(self.file_path),
                "file_size": os.path.getsize(self.file_path),
            }
            
            if front_matter:
                metadata["front_matter"] = front_matter
            
            yield Document(content=content, metadata=metadata)
        except Exception as e:
            logger.error(f"Error lazy loading Markdown file {self.file_path}: {e}")
            raise


class JSONLoader(BaseLoader):
    """
    Loader for JSON files.
    """
    
    def __init__(
        self, 
        file_path: str, 
        jq_schema: Optional[str] = None,
        text_key: Optional[str] = None,
        metadata_keys: Optional[List[str]] = None,
        array_handler: str = "concatenate",
    ):
        """
        Initialize a JSON loader.
        
        Args:
            file_path: Path to the JSON file
            jq_schema: jq-style extraction schema (if jq is available)
            text_key: Key for the text content (if not using jq)
            metadata_keys: Keys to include in metadata (if not using jq)
            array_handler: How to handle arrays: 'concatenate', 'separate', or 'flatten'
        """
        self.file_path = file_path
        self.jq_schema = jq_schema
        self.text_key = text_key
        self.metadata_keys = metadata_keys or []
        self.array_handler = array_handler
        
        if jq_schema:
            try:
                import jq
                self.jq_available = True
            except ImportError:
                logger.warning(
                    "jq-schema specified but jq package not available. "
                    "Install with `pip install jq`. Falling back to standard parsing."
                )
                self.jq_available = False
        else:
            self.jq_available = False
    
    def _process_json_with_jq(self, data: Any) -> List[Dict[str, Any]]:
        """
        Process JSON data using jq schema.
        
        Args:
            data: Parsed JSON data
            
        Returns:
            List of dictionaries with text and metadata extracted using jq
        """
        import jq
        
        # Apply jq schema to extract data
        try:
            results = jq.compile(self.jq_schema).input(data).all()
            
            # Ensure results are in the expected format
            if not results:
                logger.warning(f"jq schema '{self.jq_schema}' did not match any data")
                return []
            
            return results
        except Exception as e:
            logger.error(f"Error applying jq schema: {e}")
            raise
    
    def _process_json_standard(self, data: Any) -> List[Dict[str, Any]]:
        """
        Process JSON data using standard Python parsing.
        
        Args:
            data: Parsed JSON data
            
        Returns:
            List of dictionaries with text and metadata
        """
        results = []
        
        # Handle different data types
        if isinstance(data, dict):
            # Process single dictionary
            text = ""
            metadata = {}
            
            # Extract text content
            if self.text_key and self.text_key in data:
                text = str(data[self.text_key])
            elif not self.text_key:
                # Use entire JSON as text if no key specified
                text = json.dumps(data)
            
            # Extract metadata
            for key in self.metadata_keys:
                if key in data:
                    metadata[key] = data[key]
            
            results.append({"text": text, "metadata": metadata})
            
        elif isinstance(data, list):
            # Process list of items
            if self.array_handler == "separate":
                # Create separate documents for each item
                for item in data:
                    if isinstance(item, dict):
                        # Similar processing as above for each dictionary
                        text = ""
                        metadata = {}
                        
                        if self.text_key and self.text_key in item:
                            text = str(item[self.text_key])
                        elif not self.text_key:
                            text = json.dumps(item)
                        
                        for key in self.metadata_keys:
                            if key in item:
                                metadata[key] = item[key]
                        
                        results.append({"text": text, "metadata": metadata})
                    else:
                        # For non-dictionary items, convert to string
                        results.append({"text": str(item), "metadata": {}})
            
            elif self.array_handler == "concatenate":
                # Combine all items into one document
                all_text = ""
                all_metadata = {}
                
                for item in data:
                    if isinstance(item, dict):
                        # Extract text
                        if self.text_key and self.text_key in item:
                            all_text += str(item[self.text_key]) + "\n\n"
                        elif not self.text_key:
                            all_text += json.dumps(item) + "\n\n"
                        
                        # Extract metadata
                        for key in self.metadata_keys:
                            if key in item:
                                if key not in all_metadata:
                                    all_metadata[key] = [item[key]]
                                else:
                                    all_metadata[key].append(item[key])
                    else:
                        all_text += str(item) + "\n\n"
                
                results.append({"text": all_text, "metadata": all_metadata})
            
            elif self.array_handler == "flatten":
                # Recursively process all elements
                for item in data:
                    results.extend(self._process_json_standard(item))
        
        else:
            # Handle primitive types
            results.append({"text": str(data), "metadata": {}})
        
        return results
    
    def load(self) -> List[Document]:
        """
        Load a JSON file.
        
        Returns:
            List of Document objects extracted from the JSON
        """
        try:
            with open(self.file_path, "r") as f:
                data = json.load(f)
            
            if self.jq_available and self.jq_schema:
                extracted_items = self._process_json_with_jq(data)
            else:
                extracted_items = self._process_json_standard(data)
            
            documents = []
            file_metadata = {
                "source": self.file_path,
                "file_type": "json",
                "file_name": os.path.basename(self.file_path),
                "file_path": os.path.abspath(self.file_path),
                "file_size": os.path.getsize(self.file_path),
            }
            
            for item in extracted_items:
                if isinstance(item, dict) and "text" in item:
                    # Combined metadata from file and extracted item
                    item_metadata = {**file_metadata}
                    if "metadata" in item and isinstance(item["metadata"], dict):
                        item_metadata.update(item["metadata"])
                    
                    documents.append(Document(content=item["text"], metadata=item_metadata))
                else:
                    # If item doesn't have expected structure, convert to string
                    text = json.dumps(item) if not isinstance(item, str) else item
                    documents.append(Document(content=text, metadata=file_metadata))
            
            return documents
        except Exception as e:
            logger.error(f"Error loading JSON file {self.file_path}: {e}")
            raise
    
    def lazy_load(self) -> Iterator[Document]:
        """
        Lazily load a JSON file.
        
        Yields:
            Document objects extracted from the JSON
        """
        documents = self.load()  # For JSON, lazy loading is same as eager loading
        for doc in documents:
            yield doc


class DirectoryLoader(BaseLoader):
    """
    Loader for loading all supported files in a directory.
    """
    
    LOADER_MAPPING = {
        ".txt": TextLoader,
        ".pdf": PDFLoader,
        ".html": HTMLLoader,
        ".htm": HTMLLoader,
        ".md": MarkdownLoader,
        ".markdown": MarkdownLoader,
        ".json": JSONLoader,
    }
    
    def __init__(
        self, 
        directory_path: str,
        glob_pattern: str = "*.*",
        recursive: bool = False,
        exclude_hidden: bool = True,
        loader_mapping: Optional[Dict[str, BaseLoader]] = None,
        loader_kwargs: Optional[Dict[str, Dict[str, Any]]] = None,
        sample_size: Optional[int] = None,
        file_limit: Optional[int] = None,
    ):
        """
        Initialize a directory loader.
        
        Args:
            directory_path: Path to the directory
            glob_pattern: Glob pattern for file matching
            recursive: Whether to recursively search subdirectories
            exclude_hidden: Whether to exclude hidden files
            loader_mapping: Optional mapping of file extensions to loaders
            loader_kwargs: Optional mapping of file extensions to loader kwargs
            sample_size: Optional number of random files to sample
            file_limit: Optional maximum number of files to load
        """
        self.directory_path = directory_path
        self.glob_pattern = glob_pattern
        self.recursive = recursive
        self.exclude_hidden = exclude_hidden
        self.loader_mapping = loader_mapping or self.LOADER_MAPPING
        self.loader_kwargs = loader_kwargs or {}
        self.sample_size = sample_size
        self.file_limit = file_limit
    
    def _is_hidden(self, path: str) -> bool:
        """
        Check if a file or directory is hidden.
        
        Args:
            path: Path to check
            
        Returns:
            True if hidden, False otherwise
        """
        return os.path.basename(path).startswith(".")
    
    def _get_files(self) -> List[str]:
        """
        Get all matching files in the directory.
        
        Returns:
            List of file paths
        """
        if self.recursive:
            glob_func = Path(self.directory_path).rglob
        else:
            glob_func = Path(self.directory_path).glob
        
        files = [
            str(path) for path in glob_func(self.glob_pattern)
            if path.is_file() and (not self.exclude_hidden or not self._is_hidden(str(path)))
        ]
        
        # Apply sample_size if specified
        if self.sample_size and len(files) > self.sample_size:
            import random
            files = random.sample(files, self.sample_size)
        
        # Apply file_limit if specified
        if self.file_limit and len(files) > self.file_limit:
            files = files[:self.file_limit]
        
        return files
    
    def _get_loader_for_file(self, file_path: str) -> Optional[BaseLoader]:
        """
        Get the appropriate loader for a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Loader instance or None if no loader is available
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.loader_mapping:
            logger.warning(f"No loader available for extension {ext}")
            return None
        
        loader_class = self.loader_mapping[ext]
        kwargs = self.loader_kwargs.get(ext, {})
        
        try:
            return loader_class(file_path=file_path, **kwargs)
        except Exception as e:
            logger.error(f"Error creating loader for {file_path}: {e}")
            return None
    
    def load(self) -> List[Document]:
        """
        Load all supported files in the directory.
        
        Returns:
            List of Document objects
        """
        all_files = self._get_files()
        all_documents = []
        
        for file_path in all_files:
            loader = self._get_loader_for_file(file_path)
            if loader:
                try:
                    documents = loader.load()
                    all_documents.extend(documents)
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")
        
        return all_documents
    
    def lazy_load(self) -> Iterator[Document]:
        """
        Lazily load all supported files in the directory.
        
        Yields:
            Document objects
        """
        all_files = self._get_files()
        
        for file_path in all_files:
            loader = self._get_loader_for_file(file_path)
            if loader:
                try:
                    for document in loader.lazy_load():
                        yield document
                except Exception as e:
                    logger.error(f"Error lazy loading {file_path}: {e}")


# Add more document loaders as needed for other file types 